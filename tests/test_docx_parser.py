"""Cobre as duas falhas do DOCXParser que deixavam o índice vazio.

1. O estilo do parágrafo era lido como atributo de `<w:p>` (`w:styleId`), mas
   vive em `<w:pPr><w:pStyle w:val=...>`. O ramo de heading nunca executava.
   Além disso a comparação era contra o styleId cru, então Word em pt-BR
   (`Ttulo1`) falharia mesmo com a leitura corrigida.
2. O loop de tabelas dava `break` no primeiro item de `doc.tables`, então toda
   tabela do documento era renderizada como a primeira.

Sem heading, `md_parser` não cria nós (`^#{1,4} `) e o documento entra no RAG
como um bloco único e opaco — foi o que aconteceu com o receituário agronômico.
"""
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.parser import parse_xml

from parsers.docx_parser import DOCXParser, _heading_level, _level_from_token
from parsers.md_parser import MDParser

_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _parse(doc, tmp_path: Path) -> str:
    f = tmp_path / "d.docx"
    doc.save(str(f))
    return DOCXParser().parse(f).text


def _headings(text: str) -> list[str]:
    return [ln for ln in text.splitlines() if ln.startswith("#")]


def _set_style_id(paragraph, style_id: str) -> None:
    """Grava w:pPr/w:pStyle/@w:val direto, simulando um styleId localizado
    (pt-BR) sem depender de o estilo existir em styles.xml."""
    ppr = paragraph._p.get_or_add_pPr()
    ppr.append(parse_xml(f'<w:pStyle {_W} w:val="{style_id}"/>'))


# ── headings ──────────────────────────────────────────────────────────────────

def test_headings_ingles_viram_markdown(tmp_path):
    doc = Document()
    doc.add_heading("Primeira", level=1)
    doc.add_paragraph("corpo um")
    doc.add_heading("Segunda", level=2)
    doc.add_paragraph("corpo dois")

    assert _headings(_parse(doc, tmp_path)) == ["# Primeira", "## Segunda"]


def test_heading_ptbr_ttulo_e_reconhecido(tmp_path):
    """Regressão do bug real: Word pt-BR emite styleId `Ttulo1`, que o
    python-docx resolve para o nome canônico `Heading 1`."""
    doc = Document()
    p1 = doc.add_paragraph("Identificação do Produtor")
    _set_style_id(p1, "Ttulo1")
    doc.add_paragraph("corpo")
    p2 = doc.add_paragraph("Subitem")
    _set_style_id(p2, "Ttulo2")

    assert _headings(_parse(doc, tmp_path)) == [
        "# Identificação do Produtor",
        "## Subitem",
    ]


def test_style_id_cru_nao_e_mais_usado(tmp_path):
    """O código antigo fazia `block.get(qn('w:styleId'))`, que é sempre vazio.
    Este teste falha se alguém voltar a depender do atributo."""
    doc = Document()
    p = doc.add_paragraph("Título por estilo")
    _set_style_id(p, "Ttulo1")
    f = tmp_path / "d.docx"
    doc.save(str(f))

    reread = Document(str(f))
    block = [b for b in reread.element.body if b.tag.endswith("}p")][0]
    from docx.oxml.ns import qn
    assert block.get(qn("w:styleId")) is None, "atributo w:styleId não existe em w:p"

    assert _headings(DOCXParser().parse(f).text) == ["# Título por estilo"]


def test_heading_nivel_acima_de_4_e_limitado(tmp_path):
    """md_parser só reconhece `^#{1,4} `; um `#####` viraria texto solto."""
    doc = Document()
    p = doc.add_paragraph("Muito fundo")
    _set_style_id(p, "Ttulo7")

    heads = _headings(_parse(doc, tmp_path))
    assert heads == ["#### Muito fundo"]


def test_paragrafo_normal_nao_e_heading(tmp_path):
    doc = Document()
    doc.add_paragraph("texto comum")
    assert _headings(_parse(doc, tmp_path)) == []


@pytest.mark.parametrize(
    "token,expected",
    [
        ("Heading 1", 1),        # nome canônico
        ("heading 3", 3),
        ("Heading1", 1),         # styleId cru (sem espaço)
        ("Ttulo2", 2),           # styleId pt-BR como o Word grava
        ("Título 2", 2),
        ("Titulo 2", 2),
        ("Heading 9", 4),        # limitado ao teto
        ("Heading 0", None),
        ("Normal", None),
        ("List Paragraph", None),
        ("PargrafodaLista", None),
        ("Heading", None),       # sem número
        ("Subtitle", None),
        ("", None),
        (None, None),
    ],
)
def test_level_from_token_matriz(token, expected):
    assert _level_from_token(token) == expected


def test_heading_level_sem_estilo(tmp_path):
    """Parágrafo sem w:pPr e sem estilo resolvido não deve virar heading."""
    doc = Document()
    doc.add_paragraph("solto")
    f = _save(doc, tmp_path)
    reread = Document(str(f))
    assert _heading_level(reread.paragraphs[0]) is None


# ── tabelas ───────────────────────────────────────────────────────────────────

def test_cada_tabela_entra_uma_vez_e_na_ordem(tmp_path):
    """Regressão do `break`: antes as 3 tabelas saíam como 3 cópias da primeira."""
    doc = Document()
    for i in range(3):
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = f"H{i}A"
        t.cell(0, 1).text = f"H{i}B"
        t.cell(1, 0).text = f"linha-unica-{i}"
        t.cell(1, 1).text = f"valor-{i}"

    text = _parse(doc, tmp_path)

    for i in range(3):
        assert text.count(f"linha-unica-{i}") == 1, f"tabela {i} ausente ou duplicada"
    # ordem preservada
    assert text.index("linha-unica-0") < text.index("linha-unica-1") < text.index("linha-unica-2")
    # um separador de header por tabela
    assert sum(1 for ln in text.splitlines() if set(ln.replace("|", "").strip()) == {"-"}) == 3


def test_todas_as_linhas_de_tabela_sao_emitidas(tmp_path):
    doc = Document()
    t = doc.add_table(rows=4, cols=3)
    for r in range(4):
        for c in range(3):
            t.cell(r, c).text = f"r{r}c{c}"

    text = _parse(doc, tmp_path)
    data_rows = [
        ln for ln in text.splitlines()
        if ln.startswith("| ") and set(ln.replace("|", "").strip()) != {"-"}
    ]
    assert len(data_rows) == 4


# ── extração de texto ─────────────────────────────────────────────────────────

def test_texto_de_caixa_de_texto_e_preservado(tmp_path):
    """`Paragraph.text` ignora w:txbxContent. O parser varre todos os <w:t>
    descendentes justamente para não perder avisos em caixas de texto."""
    doc = Document()
    p = doc.add_paragraph("corpo ")
    p._p.append(parse_xml(
        f'<w:r {_W} xmlns:v="urn:schemas-microsoft-com:vml"><w:pict><v:shape><v:textbox>'
        '<w:txbxContent><w:p><w:r><w:t>AVISO-NA-CAIXA</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict></w:r>'
    ))

    assert "AVISO-NA-CAIXA" in _parse(doc, tmp_path)


def test_hyperlink_e_preservado(tmp_path):
    doc = Document()
    p = doc.add_paragraph("veja ")
    p._p.append(parse_xml(
        f'<w:hyperlink {_W} xmlns:r="http://schemas.openxmlformats.org/'
        'officeDocument/2006/relationships" r:id="rId1">'
        '<w:r><w:t>TEXTO-DO-LINK</w:t></w:r></w:hyperlink>'
    ))

    assert "TEXTO-DO-LINK" in _parse(doc, tmp_path)


def test_paragrafos_vazios_sao_ignorados(tmp_path):
    doc = Document()
    doc.add_heading("Titulo", level=1)
    doc.add_paragraph("   ")
    doc.add_paragraph("")
    doc.add_paragraph("conteudo")

    text = _parse(doc, tmp_path)
    assert [ln for ln in text.splitlines() if ln.strip()] == ["# Titulo", "conteudo"]


# ── integração com o md_parser (o que de fato alimenta o índice) ───────────────

def test_secoes_geram_nos_no_md_parser(tmp_path):
    """O ponto final: heading no DOCX -> nó no índice do PageIndex."""
    doc = Document()
    for i in range(1, 4):
        doc.add_heading(f"Secao {i}", level=1)
        doc.add_paragraph(f"conteudo da secao {i}")

    parsed = DOCXParser().parse(_save(doc, tmp_path))
    md = tmp_path / "out.md"
    md.write_text(parsed.text, encoding="utf-8")

    page_map = MDParser().parse(md).page_map
    assert [p.label for p in page_map] == ["Secao 1", "Secao 2", "Secao 3"]


def test_documento_sem_heading_gera_no_unico(tmp_path):
    """Limitação conhecida: DOCX cujos títulos são formatação manual (só estilo
    `Normal`) não tem como ser fatiado — vira um nó só. Documentado aqui para
    que a diferença seja intencional, não surpresa."""
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"paragrafo {i}")

    parsed = DOCXParser().parse(_save(doc, tmp_path))
    md = tmp_path / "out.md"
    md.write_text(parsed.text, encoding="utf-8")

    page_map = MDParser().parse(md).page_map
    assert len(page_map) == 1
    assert page_map[0].label == "Documento"


def test_page_map_conta_secoes(tmp_path):
    doc = Document()
    doc.add_paragraph("preambulo")
    doc.add_heading("A", level=1)
    doc.add_paragraph("a")
    doc.add_heading("B", level=1)
    doc.add_paragraph("b")

    parsed = DOCXParser().parse(_save(doc, tmp_path))
    # preambulo + A + B
    assert len(parsed.page_map) == 3


def _save(doc, tmp_path: Path) -> Path:
    f = tmp_path / "d.docx"
    doc.save(str(f))
    return f
