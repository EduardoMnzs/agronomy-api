import re
from pathlib import Path

# Side-effect: registra parser XML seguro antes do python-docx (anti billion-laughs/XXE).
import defusedxml.ElementTree  # noqa: F401

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from parsers.base import BaseParser, PageRef, ParsedDocument
from parsers.safety import assert_zip_safe

# `md_parser` só reconhece headings de H1 a H4 (`^#{1,4} `). Um `#####` não vira
# nó e o conteúdo é absorvido pelo bloco anterior, então limitamos aqui.
_MAX_HEADING_LEVEL = 4

# Casa o nome canônico ("Heading 2") e o localizado, com ou sem espaço, porque a
# mesma regex serve para o nome resolvido e para o styleId cru.
# O `[íi]?` é essencial: o Word DESCARTA caracteres não-ASCII ao gerar styleId,
# então "Título 2" vira `Ttulo2` — sem o "í", não substituído por "i".
_HEADING_TOKEN_RE = re.compile(r"^(?:heading|t[íi]?tulo)\s*(\d+)$", re.IGNORECASE)


def _level_from_token(token: str | None) -> int | None:
    """Extrai o nível de um rótulo de estilo, limitado a _MAX_HEADING_LEVEL."""
    match = _HEADING_TOKEN_RE.match((token or "").strip())
    if not match:
        return None
    level = int(match.group(1))
    if level < 1:
        return None
    return min(level, _MAX_HEADING_LEVEL)


def _raw_style_id(paragraph: Paragraph) -> str | None:
    """styleId cru em w:pPr/w:pStyle/@w:val, sem passar pela resolução de estilo."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return None
    return p_style.get(qn("w:val"))


def _heading_level(paragraph: Paragraph) -> int | None:
    """Nível do heading (1.._MAX_HEADING_LEVEL) do parágrafo, ou None.

    Duas fontes, nessa ordem:

    1. O NOME resolvido do estilo. É o caminho normal: em documentos pt-BR o
       styleId é `Ttulo1` e o python-docx resolve para `Heading 1` usando a
       definição builtin de `styles.xml`. Comparar contra o id cru — como fazia
       a versão anterior — deixa passar batido o heading de qualquer Word não
       inglês.
    2. O styleId cru, como fallback. Se o documento referencia um estilo que não
       está definido em `styles.xml` (acontece com DOCX gerado por ferramenta de
       terceiros), o python-docx cai silenciosamente para `Normal` e o passo 1
       não dá sinal nenhum. Sem esse fallback o documento voltaria a indexar
       como um bloco único, que é justamente a falha silenciosa que queremos
       impossibilitar.
    """
    style = paragraph.style
    if style is not None:
        level = _level_from_token(style.name)
        if level is not None:
            return level
    return _level_from_token(_raw_style_id(paragraph))


def _table_to_markdown(table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


class DOCXParser(BaseParser):
    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        assert_zip_safe(file_path)
        doc = Document(str(file_path))
        sections: list[str] = []
        page_map: list[PageRef] = []
        current_section: list[str] = []
        section_number = 1

        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

            if tag == "p":
                # Varre TODOS os <w:t> descendentes, não só os runs diretos do
                # parágrafo: `Paragraph.text` ignora texto dentro de caixas de
                # texto (w:txbxContent), comum em avisos e notas laterais.
                para_text = "".join(
                    node.text or "" for node in block.iter() if node.tag.endswith("}t")
                ).strip()
                if not para_text:
                    continue

                level = _heading_level(Paragraph(block, doc))
                if level is not None:
                    if current_section:
                        sections.append("\n".join(current_section))
                        page_map.append(PageRef(page_number=section_number, label=f"Seção {section_number}"))
                        section_number += 1
                        current_section = []
                    current_section.append("#" * level + " " + para_text)
                else:
                    current_section.append(para_text)

            elif tag == "tbl":
                # Constrói a Table a partir DESTE elemento. Antes o código varria
                # `doc.tables` e dava `break` no primeiro, então toda tabela do
                # documento era renderizada como a primeira — as demais nunca
                # entravam no índice.
                current_section.append(_table_to_markdown(Table(block, doc)))

        if current_section:
            sections.append("\n".join(current_section))
            page_map.append(PageRef(page_number=section_number, label=f"Seção {section_number}"))

        return ParsedDocument(
            text="\n\n".join(sections),
            page_map=page_map,
            doc_type="docx",
            original_filename=file_path.name,
        )
